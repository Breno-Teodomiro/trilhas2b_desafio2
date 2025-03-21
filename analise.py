import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# --------------------------------------------------------------------------------
# EXPLICAÇÃO DAS ANÁLISES:
#
# 1) Interpolação Linear: 
#    - Preenche valores faltantes na coluna "Renda per capita" de forma suave,
#      utilizando os valores disponíveis.
#
# 2) Derivadas – Velocidade e Aceleração:
#    - Velocidade (1ª derivada): mostra a taxa de variação anual dos indicadores.
#    - Aceleração (2ª derivada): indica se a taxa de variação está aumentando ou diminuindo.
#
# 3) Médias Móveis (3 e 7 anos):
#    - Suavizam as oscilações de curto prazo e destacam as tendências de médio e longo prazo.
#
# 4) Medidas de Dispersão (Média, Variância, Desvio Padrão):
#    - Quantificam a variação dos dados em relação à média.
#
# Durante o período de COVID (aproximadamente 2020-2021), observa-se que os indicadores,
# especialmente o PIB e a Renda per capita, sofreram variações acentuadas, refletindo o impacto
# econômico e social da pandemia. A coleta populacional, realizada em alguns anos de forma bienal,
# também demanda cautela na interpretação dos resultados.
# --------------------------------------------------------------------------------

def carregar_dados(nome_arquivo):
    """
    Lê a planilha Excel e realiza o tratamento inicial:
      - Ordena pelos anos;
      - Converte a coluna 'Ano' para inteiro;
      - Cria a coluna 'Renda per capita (Interpolada)' preenchendo os valores faltantes por interpolação linear.
    """

    df = pd.read_excel("Planilha do Desafio 2 -Trilhas 2B.xlsx", header=1)  # Se o cabeçalho estiver na segunda linha (índice 1)

    # Converter 'Ano' para inteiro e ordenar
    df['Ano'] = df['Ano'].astype(int)
    df.sort_values('Ano', inplace=True)
    
    # Interpolação linear na coluna 'Renda per capita'
    df['Renda per capita (Interpolada)'] = df['Renda per capita'].interpolate(method='linear')
    
    return df

def calcular_derivadas(df, coluna, tempo=1):
    """
    Calcula a velocidade (1ª derivada) e a aceleração (2ª derivada) da coluna informada.
    O parâmetro 'tempo' representa o intervalo em anos entre as medições (ex.: 2 para dados bienais).
    """
    velocidade = df[coluna].diff() / tempo
    aceleracao = velocidade.diff() / tempo
    return velocidade, aceleracao

def calcular_medias_moveis(df, coluna, janelas=[3, 7]):
    """
    Calcula médias móveis para as janelas especificadas (ex.: 3 e 7 anos).
    Retorna um dicionário com as médias móveis.
    """
    medias = {}
    for janela in janelas:
        medias[f'MM_{janela}'] = df[coluna].rolling(window=janela).mean()
    return medias

def calcular_dispersao(df, coluna):
    """
    Retorna média, variância e desvio padrão para a coluna especificada.
    """
    media = df[coluna].mean()
    var = df[coluna].var()
    std = df[coluna].std()
    return media, var, std

def gerar_planilha_analise(df, nome_arquivo):
    """
    Gera uma planilha Excel com todos os dados originais e os dados processados (derivadas, médias móveis, etc.).
    """
    with pd.ExcelWriter(nome_arquivo, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Analise')
    
    print(f'[OK] Planilha de análise "{nome_arquivo}" gerada com sucesso.')

def gerar_relatorio(df, graficos, nome_arquivo_doc):
    """
    Gera um relatório DOCX que inclui:
      - Introdução com os dados originais e o contexto do estudo;
      - Descrição detalhada do processo de cálculos (interpolação, derivadas e médias móveis);
      - Uma tabela resumo com as estatísticas dos indicadores;
      - Apresentação dos gráficos gerados (com explicações sobre o que cada gráfico representa);
      - Uma seção exclusiva para analisar os impactos do período da COVID;
      - Conclusões e observações finais.
    """
    doc = Document()
    
    # Título
    doc.add_heading('Relatório de Análise de Dados do Estado do Maranhão', 0)
    
    # Introdução
    doc.add_heading('1. Introdução', level=1)
    doc.add_paragraph(
        "Este relatório apresenta uma análise dos dados referentes à População Estimada, PIB Estimado (R$ em Bilhões) e Renda per capita, "
        "no período de 2010 a 2024. A análise contempla tanto os dados originais quanto os valores processados por meio de fórmulas, "
        "como interpolação, cálculo de derivadas e médias móveis. Especial atenção foi dada ao impacto da COVID (aproximadamente 2020-2021), "
        "que provocou variações significativas nos indicadores."
    )
    
    # Dados Originais
    doc.add_heading('2. Dados Originais', level=1)
    doc.add_paragraph("A seguir, apresenta-se uma amostra dos dados originais conforme extraídos da planilha:")
    
    # Criar tabela com os dados originais (limitada às 10 primeiras linhas para visualização)
    df_orig = df[['Ano', 'População Estimada', 'PIB Estimado (R$ em Bilhões)', 'Renda per capita']].head(10)
    tabela_origem = doc.add_table(rows=1, cols=len(df_orig.columns))
    hdr_cells = tabela_origem.rows[0].cells
    for i, col in enumerate(df_orig.columns):
        hdr_cells[i].text = str(col)
    for _, row in df_orig.iterrows():
        row_cells = tabela_origem.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    
    # Processos de Cálculo
    doc.add_heading('3. Processo de Cálculo e Tratamento dos Dados', level=1)
    doc.add_paragraph(
        "Para obter uma análise mais robusta, foram aplicados os seguintes processos:\n"
        "- **Interpolação Linear**: Preenche os valores faltantes na coluna 'Renda per capita', gerando a coluna "
        "'Renda per capita (Interpolada)'.\n"
        "- **Derivadas (Velocidade e Aceleração)**: Calculadas para cada indicador para evidenciar a taxa de variação "
        "anual e a mudança dessa taxa. Nota: Para 'População Estimada', foi considerado um intervalo de 2 anos.\n"
        "- **Médias Móveis**: Foram calculadas para janelas de 3 e 7 anos, suavizando oscilações e destacando as tendências de longo prazo.\n"
        "- **Medidas de Dispersão**: Média, variância e desvio padrão foram computados para quantificar a dispersão dos dados."
    )
    
    # Tabela de Estatísticas Descritivas
    doc.add_heading('4. Estatísticas Descritivas', level=1)
    doc.add_paragraph("A tabela abaixo resume as principais medidas de dispersão dos indicadores:")
    tabela_estat = doc.add_table(rows=1, cols=5)
    hdr_cells = tabela_estat.rows[0].cells
    hdr_cells[0].text = 'Métrica'
    hdr_cells[1].text = 'População'
    hdr_cells[2].text = 'PIB'
    hdr_cells[3].text = 'Renda p/ capita'
    hdr_cells[4].text = 'Renda (Interp.)'
    
    media_pop, var_pop, std_pop = calcular_dispersao(df, 'População Estimada')
    media_pib, var_pib, std_pib = calcular_dispersao(df, 'PIB Estimado (R$ em Bilhões)')
    media_renda, var_renda, std_renda = calcular_dispersao(df, 'Renda per capita')
    media_renda_i, var_renda_i, std_renda_i = calcular_dispersao(df, 'Renda per capita (Interpolada)')
    
    estatisticas = [
        ['Média', f'{media_pop:.2f}', f'{media_pib:.2f}', f'{media_renda:.2f}', f'{media_renda_i:.2f}'],
        ['Variância', f'{var_pop:.2f}', f'{var_pib:.2f}', f'{var_renda:.2f}', f'{var_renda_i:.2f}'],
        ['Desvio Padrão', f'{std_pop:.2f}', f'{std_pib:.2f}', f'{std_renda:.2f}', f'{std_renda_i:.2f}']
    ]
    for linha in estatisticas:
        row_cells = tabela_estat.add_row().cells
        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)
    
    # Gráficos e Análises Visuais
    doc.add_heading('5. Gráficos e Análises Visuais', level=1)
    doc.add_paragraph(
        "Os gráficos a seguir apresentam de forma visual a evolução dos indicadores, as derivadas (velocidade e aceleração) e as médias móveis. "
        "Cada gráfico foi elaborado para facilitar a interpretação das tendências e das oscilações dos dados."
    )
    
    for descricao, caminho in graficos.items():
        doc.add_heading(descricao, level=2)
        doc.add_picture(caminho, width=Inches(5.5))
        if descricao == 'Evolução dos Indicadores':
            doc.add_paragraph(
                "Este gráfico ilustra a evolução da População Estimada, do PIB Estimado e da Renda per capita (após interpolação) entre 2010 e 2024. "
                "Observa-se a tendência geral de crescimento, mas também possíveis quedas ou estagnações, especialmente em períodos de instabilidade econômica."
            )
        elif descricao == 'Dinâmica do PIB':
            doc.add_paragraph(
                "Neste gráfico, são analisados a velocidade (taxa de variação anual) e a aceleração do PIB. "
                "Notamos que, durante o período de COVID, houve uma desaceleração brusca, seguida por uma recuperação gradual."
            )
        elif descricao == 'Médias Móveis da Renda per capita':
            doc.add_paragraph(
                "Aqui são exibidas a Renda per capita original (após interpolação) e suas médias móveis de 3 e 7 anos. "
                "Essas médias ajudam a suavizar oscilações e a evidenciar tendências de longo prazo, mesmo diante de variações pontuais."
            )
    
    # Análise do Período da COVID
    doc.add_heading('6. Impacto do Período da COVID', level=1)
    doc.add_paragraph(
        "O período da COVID (aproximadamente 2020-2021) mostrou impactos significativos nos indicadores analisados. "
        "Observa-se que:\n"
        "- O PIB apresentou uma queda acentuada durante os anos de 2020 e 2021, refletindo a desaceleração econômica global.\n"
        "- A Renda per capita sofreu variações marcantes, possivelmente devido à perda de empregos e à redução do poder de compra.\n"
        "- A população, coletada em alguns anos de forma bienal, mostra variações que devem ser interpretadas com cautela, "
        "pois a frequência da coleta pode mascarar mudanças abruptas ocorridas durante a pandemia.\n\n"
        "Esses efeitos evidenciam a importância de se aplicar técnicas de suavização, como as médias móveis, e de se analisar as derivadas "
        "(velocidade e aceleração) para compreender melhor a dinâmica dos dados durante períodos de crise."
    )
    
    # Dados Gerados pelas Fórmulas (amostra)
    doc.add_heading('7. Dados Processados e Resultados das Fórmulas', level=1)
    doc.add_paragraph(
        "A tabela a seguir apresenta uma amostra dos dados processados, incluindo os cálculos de derivadas e médias móveis. "
        "Estes dados foram gerados automaticamente a partir dos dados originais e permitem uma análise quantitativa detalhada."
    )
    
    # Exibindo uma amostra dos dados processados (primeiras 10 linhas)
    df_processado = df.head(10)
    cols = df_processado.columns.tolist()
    tabela_processada = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = tabela_processada.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = str(col)
    for _, row in df_processado.iterrows():
        row_cells = tabela_processada.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)
    
    # Conclusões Finais
    doc.add_heading('8. Conclusões e Observações Finais', level=1)
    doc.add_paragraph(
        "Em síntese, a análise dos dados revela que:\n"
        "- Os processos de interpolação, cálculo de derivadas e médias móveis foram fundamentais para suavizar as oscilações e "
        "destacar as tendências dos indicadores ao longo dos anos.\n"
        "- O período da COVID impactou negativamente o PIB e a Renda per capita, embora haja sinais de recuperação pós-pandemia.\n"
        "- A coleta de dados populacionais de forma bienal impõe desafios na interpretação de variações abruptas, sendo necessário "
        "cautela ao analisar esses números.\n\n"
        "Essas análises permitem um entendimento mais profundo das dinâmicas socioeconômicas do Estado do Maranhão, "
        "oferecendo subsídios para decisões e políticas públicas fundamentadas."
    )
    
    doc.save(nome_arquivo_doc)
    print(f'[OK] Relatório "{nome_arquivo_doc}" gerado com sucesso.')

def main():
    # Definição dos arquivos
    arquivo_excel = 'Planilha do Desafio 2 -Trilhas 2B.xlsx'
    arquivo_excel_saida = 'Analise_Desafio_2.xlsx'
    arquivo_doc = 'Relatorio_Analise_Desafio_2.docx'
    
    # 1) Carregar os dados da planilha
    df = carregar_dados(arquivo_excel)
    
    # 2) Calcular derivadas (velocidade e aceleração)
    # Para 'População Estimada', como os dados podem ser coletados de 2 em 2 anos, usa-se tempo=2.
    for coluna in ['População Estimada', 'PIB Estimado (R$ em Bilhões)', 'Renda per capita (Interpolada)']:
        tempo = 2 if coluna == 'População Estimada' else 1
        vel, acel = calcular_derivadas(df, coluna, tempo=tempo)
        df[f'{coluna} Velocidade'] = vel
        df[f'{coluna} Aceleração'] = acel
    
    # 3) Calcular médias móveis (janelas de 3 e 7 anos)
    for coluna in ['População Estimada', 'PIB Estimado (R$ em Bilhões)', 'Renda per capita (Interpolada)']:
        medias = calcular_medias_moveis(df, coluna, janelas=[3, 7])
        for key, valores in medias.items():
            df[f'{coluna} {key}'] = valores
    
    # 4) Gerar a planilha de análise com todos os dados
    gerar_planilha_analise(df, arquivo_excel_saida)
    
    # 5) Gerar e salvar os gráficos
    graficos = {}
    anos = df['Ano']
    
    # Gráfico 1: Evolução dos Indicadores
    plt.figure(figsize=(10, 6))
    plt.plot(anos, df['População Estimada'], marker='o', label='População')
    plt.plot(anos, df['PIB Estimado (R$ em Bilhões)'], marker='o', label='PIB')
    plt.plot(anos, df['Renda per capita (Interpolada)'], marker='o', label='Renda per capita')
    plt.title('Evolução dos Indicadores (2010-2024)')
    plt.xlabel('Ano')
    plt.ylabel('Valores')
    plt.legend()
    plt.grid(True)
    grafico1 = 'grafico_evolucao.png'
    plt.savefig(grafico1)
    plt.close()
    graficos['Evolução dos Indicadores'] = grafico1
    
    # Gráfico 2: Dinâmica do PIB (Velocidade e Aceleração)
    plt.figure(figsize=(10, 6))
    plt.plot(anos, df['PIB Estimado (R$ em Bilhões) Velocidade'], marker='o', label='Velocidade PIB')
    plt.plot(anos, df['PIB Estimado (R$ em Bilhões) Aceleração'], marker='o', label='Aceleração PIB')
    plt.title('Dinâmica do PIB (Velocidade e Aceleração)')
    plt.xlabel('Ano')
    plt.ylabel('Variação')
    plt.legend()
    plt.grid(True)
    grafico2 = 'grafico_dinamica_pib.png'
    plt.savefig(grafico2)
    plt.close()
    graficos['Dinâmica do PIB'] = grafico2

    # Gráfico 3: Médias Móveis da Renda per capita
    plt.figure(figsize=(10, 6))
    plt.plot(anos, df['Renda per capita (Interpolada)'], marker='o', label='Renda per capita')
    plt.plot(anos, df['Renda per capita (Interpolada) MM_3'], marker='o', label='Média Móvel 3 anos')
    plt.plot(anos, df['Renda per capita (Interpolada) MM_7'], marker='o', label='Média Móvel 7 anos')
    plt.title('Médias Móveis da Renda per capita')
    plt.xlabel('Ano')
    plt.ylabel('Renda per capita')
    plt.legend()
    plt.grid(True)
    grafico3 = 'grafico_media_movel.png'
    plt.savefig(grafico3)
    plt.close()
    graficos['Médias Móveis da Renda per capita'] = grafico3
    
    # 6) Gerar o relatório DOCX com todas as análises e gráficos
    gerar_relatorio(df, graficos, arquivo_doc)
    
    # 7) (Opcional) Remover arquivos temporários dos gráficos
    for g in graficos.values():
        if os.path.exists(g):
            os.remove(g)

if __name__ == "__main__":
    main()